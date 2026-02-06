from flask import Flask, render_template, request, send_file, jsonify
from flask_sqlalchemy import SQLAlchemy
import pdfplumber
from docx import Document
import io
import os

app = Flask(__name__)

# Podešavanje baze za produkciju (Render/Cloud)
basedir = os.path.abspath(os.path.dirname(__file__))
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///' + os.path.join(basedir, 'pdf_pro.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)

# Model baze za istoriju obrada
class History(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(100))
    format = db.Column(db.String(10))
    pages = db.Column(db.Integer)

# Inicijalizacija baze
with app.app_context():
    db.create_all()

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/stats', methods=['GET'])
def get_stats():
    try:
        h = History.query.order_by(History.id.desc()).limit(10).all()
        total_pages = db.session.query(db.func.sum(History.pages)).scalar() or 0
        items = [{"id": i.id, "name": i.filename, "fmt": i.format} for i in h]
        return jsonify({"total_pages": total_pages, "history": items})
    except Exception as e:
        return jsonify({"total_pages": 0, "history": [], "error": str(e)})

@app.route('/api/export', methods=['POST'])
def export_file():
    if 'file' not in request.files:
        return "Niste poslali fajl", 400
    
    file = request.files['file']
    fmt = request.form.get('format', 'txt')
    
    if file.filename == '':
        return "Nije izabran fajl", 400

    text_content = ""
    
    try:
        # Čitanje PDF-a sa pdfplumberom (podrška za ćirilicu/latinicu)
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                # x_tolerance pomaže kod spojenih slova
                val = page.extract_text(x_tolerance=2)
                if val:
                    text_content += val + "\n\n"
            num_pages = len(pdf.pages)

        if not text_content.strip():
            text_content = "PAŽNJA: Tekst nije pronađen. Moguće je da je PDF skeniran kao slika."

        # Čuvanje u bazu
        new_entry = History(filename=file.filename, format=fmt, pages=num_pages)
        db.session.add(new_entry)
        db.session.commit()

        output = io.BytesIO()
        
        if fmt == 'docx':
            doc = Document()
            doc.add_heading(f"Eksportovano iz PDF: {file.filename}", 0)
            doc.add_paragraph(text_content)
            doc.save(output)
            download_name = f"{os.path.splitext(file.filename)[0]}.docx"
            mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            output.write(text_content.encode('utf-8'))
            download_name = f"{os.path.splitext(file.filename)[0]}.txt"
            mimetype = "text/plain"
        
        output.seek(0)
        return send_file(
            output, 
            as_attachment=True, 
            download_name=download_name, 
            mimetype=mimetype
        )
        
    except Exception as e:
        print(f"Greska: {e}")
        return f"Došlo je do greške: {str(e)}", 500

if __name__ == '__main__':
    # Uzimanje porta koji server (Render) dodeli, inače koristi 1010
    port = int(os.environ.get("PORT", 1010))
    # Na serveru host MORA biti 0.0.0.0
    app.run(host='0.0.0.0', port=port)